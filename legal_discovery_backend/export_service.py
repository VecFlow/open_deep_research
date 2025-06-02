"""
Export Service for Legal Discovery Backend.
Handles export of analysis results to Word documents and PDFs.
"""

import io
import logging
from datetime import datetime
from typing import Tuple, Dict, Any, List

# Document generation imports
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from .models import AnalysisDB, AnalysisCategory
except ImportError:
    from models import AnalysisDB, AnalysisCategory

logger = logging.getLogger(__name__)

class ExportService:
    """Service for exporting analysis results to various formats."""
    
    def __init__(self):
        pass
    
    async def export_to_word(self, analysis_db: AnalysisDB) -> Tuple[bytes, str]:
        """Export analysis results to Word document."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for Word export")
        
        try:
            # Create new document
            doc = Document()
            
            # Add title
            title = doc.add_heading('Legal Analysis Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add case information
            doc.add_heading('Case Information', level=1)
            
            case_info_table = doc.add_table(rows=6, cols=2)
            case_info_table.style = 'Table Grid'
            
            info_rows = [
                ('Analysis ID:', analysis_db.id),
                ('Case ID:', analysis_db.case_id),
                ('Status:', analysis_db.status.title()),
                ('Created:', analysis_db.created_at.strftime('%Y-%m-%d %H:%M:%S')),
                ('Completed:', analysis_db.completed_at.strftime('%Y-%m-%d %H:%M:%S') if analysis_db.completed_at else 'In Progress'),
                ('Progress:', f'{analysis_db.progress_percentage}%')
            ]
            
            for i, (label, value) in enumerate(info_rows):
                case_info_table.cell(i, 0).text = label
                case_info_table.cell(i, 1).text = str(value)
            
            doc.add_paragraph()
            
            # Add analysis categories
            if analysis_db.categories:
                doc.add_heading('Analysis Categories', level=1)
                
                for i, category_data in enumerate(analysis_db.categories, 1):
                    category = AnalysisCategory(**category_data) if isinstance(category_data, dict) else category_data
                    
                    doc.add_heading(f'{i}. {category.name}', level=2)
                    doc.add_paragraph(f'Description: {category.description}')
                    
                    if category.content:
                        doc.add_paragraph('Analysis:')
                        content_para = doc.add_paragraph(category.content)
                        content_para.style = 'Quote'
                    
                    doc.add_paragraph()
            
            # Add completed categories
            if analysis_db.completed_categories:
                doc.add_heading('Completed Analysis', level=1)
                
                for i, category_data in enumerate(analysis_db.completed_categories, 1):
                    category = AnalysisCategory(**category_data) if isinstance(category_data, dict) else category_data
                    
                    doc.add_heading(f'{i}. {category.name}', level=2)
                    doc.add_paragraph(f'Description: {category.description}')
                    
                    if category.content:
                        doc.add_paragraph('Analysis Results:')
                        content_para = doc.add_paragraph(category.content)
                        content_para.style = 'Quote'
                    
                    doc.add_paragraph()
            
            # Add deposition questions
            if analysis_db.deposition_questions:
                doc.add_heading('Deposition Questions', level=1)
                
                deposition_data = analysis_db.deposition_questions
                if isinstance(deposition_data, dict):
                    witness_questions = deposition_data.get('witness_questions', [])
                    
                    for witness_data in witness_questions:
                        doc.add_heading(f"Witness: {witness_data.get('witness_name', 'Unknown')}", level=2)
                        doc.add_paragraph(f"Role: {witness_data.get('witness_role', 'Not specified')}")
                        
                        questions = witness_data.get('questions', [])
                        for i, question_data in enumerate(questions, 1):
                            doc.add_paragraph(f"Q{i}: {question_data.get('question', '')}")
                            doc.add_paragraph(f"Purpose: {question_data.get('purpose', '')}")
                            
                            expected_areas = question_data.get('expected_areas', [])
                            if expected_areas:
                                doc.add_paragraph(f"Expected Areas: {', '.join(expected_areas)}")
                            
                            doc.add_paragraph()
            
            # Add final analysis
            if analysis_db.final_analysis:
                doc.add_heading('Final Analysis', level=1)
                final_para = doc.add_paragraph(analysis_db.final_analysis)
                final_para.style = 'Quote'
            
            # Add footer
            doc.add_paragraph()
            footer_para = doc.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_run = footer_para.add_run(f'Generated on {datetime.now().strftime("%Y-%m-%d %H:%M:%S")} by Legal Discovery System')
            footer_run.italic = True
            
            # Save to bytes
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            # Generate filename
            safe_case_id = analysis_db.case_id.replace('/', '_').replace('\\', '_')
            filename = f'legal_analysis_{safe_case_id}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
            
            return doc_buffer.getvalue(), filename
            
        except Exception as e:
            logger.error(f"Failed to export analysis {analysis_db.id} to Word: {e}")
            raise
    
    async def export_to_pdf(self, analysis_db: AnalysisDB) -> Tuple[bytes, str]:
        """Export analysis results to PDF document."""
        if not PDF_AVAILABLE:
            raise ImportError("reportlab is required for PDF export")
        
        try:
            # Create PDF buffer
            pdf_buffer = io.BytesIO()
            
            # Create document
            doc = SimpleDocTemplate(
                pdf_buffer,
                pagesize=letter,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=18
            )
            
            # Get styles
            styles = getSampleStyleSheet()
            
            # Create custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
                alignment=1  # Center alignment
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=16,
                spaceBefore=20,
                spaceAfter=12
            )
            
            subheading_style = ParagraphStyle(
                'CustomSubheading',
                parent=styles['Heading3'],
                fontSize=14,
                spaceBefore=16,
                spaceAfter=8
            )
            
            # Build document content
            story = []
            
            # Title
            story.append(Paragraph('Legal Analysis Report', title_style))
            story.append(Spacer(1, 12))
            
            # Case information table
            story.append(Paragraph('Case Information', heading_style))
            
            case_info_data = [
                ['Analysis ID:', analysis_db.id],
                ['Case ID:', analysis_db.case_id],
                ['Status:', analysis_db.status.title()],
                ['Created:', analysis_db.created_at.strftime('%Y-%m-%d %H:%M:%S')],
                ['Completed:', analysis_db.completed_at.strftime('%Y-%m-%d %H:%M:%S') if analysis_db.completed_at else 'In Progress'],
                ['Progress:', f'{analysis_db.progress_percentage}%']
            ]
            
            case_info_table = Table(case_info_data, colWidths=[2*inch, 4*inch])
            case_info_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(case_info_table)
            story.append(Spacer(1, 20))
            
            # Analysis categories
            if analysis_db.categories:
                story.append(Paragraph('Analysis Categories', heading_style))
                
                for i, category_data in enumerate(analysis_db.categories, 1):
                    category = AnalysisCategory(**category_data) if isinstance(category_data, dict) else category_data
                    
                    story.append(Paragraph(f'{i}. {category.name}', subheading_style))
                    story.append(Paragraph(f'<b>Description:</b> {category.description}', styles['Normal']))
                    
                    if category.content:
                        story.append(Paragraph('<b>Analysis:</b>', styles['Normal']))
                        story.append(Paragraph(category.content, styles['Quote']))
                    
                    story.append(Spacer(1, 12))
            
            # Completed categories
            if analysis_db.completed_categories:
                story.append(Paragraph('Completed Analysis', heading_style))
                
                for i, category_data in enumerate(analysis_db.completed_categories, 1):
                    category = AnalysisCategory(**category_data) if isinstance(category_data, dict) else category_data
                    
                    story.append(Paragraph(f'{i}. {category.name}', subheading_style))
                    story.append(Paragraph(f'<b>Description:</b> {category.description}', styles['Normal']))
                    
                    if category.content:
                        story.append(Paragraph('<b>Analysis Results:</b>', styles['Normal']))
                        story.append(Paragraph(category.content, styles['Quote']))
                    
                    story.append(Spacer(1, 12))
            
            # Deposition questions
            if analysis_db.deposition_questions:
                story.append(Paragraph('Deposition Questions', heading_style))
                
                deposition_data = analysis_db.deposition_questions
                if isinstance(deposition_data, dict):
                    witness_questions = deposition_data.get('witness_questions', [])
                    
                    for witness_data in witness_questions:
                        story.append(Paragraph(f"<b>Witness:</b> {witness_data.get('witness_name', 'Unknown')}", subheading_style))
                        story.append(Paragraph(f"<b>Role:</b> {witness_data.get('witness_role', 'Not specified')}", styles['Normal']))
                        
                        questions = witness_data.get('questions', [])
                        for i, question_data in enumerate(questions, 1):
                            story.append(Paragraph(f"<b>Q{i}:</b> {question_data.get('question', '')}", styles['Normal']))
                            story.append(Paragraph(f"<b>Purpose:</b> {question_data.get('purpose', '')}", styles['Normal']))
                            
                            expected_areas = question_data.get('expected_areas', [])
                            if expected_areas:
                                story.append(Paragraph(f"<b>Expected Areas:</b> {', '.join(expected_areas)}", styles['Normal']))
                            
                            story.append(Spacer(1, 8))
            
            # Final analysis
            if analysis_db.final_analysis:
                story.append(Paragraph('Final Analysis', heading_style))
                story.append(Paragraph(analysis_db.final_analysis, styles['Quote']))
            
            # Footer
            story.append(Spacer(1, 30))
            footer_text = f'Generated on {datetime.now().strftime("%Y-%m-%d %H:%M:%S")} by Legal Discovery System'
            footer_style = ParagraphStyle(
                'Footer',
                parent=styles['Normal'],
                fontSize=10,
                alignment=1,  # Center alignment
                textColor=colors.grey
            )
            story.append(Paragraph(footer_text, footer_style))
            
            # Build PDF
            doc.build(story)
            
            # Get PDF bytes
            pdf_buffer.seek(0)
            
            # Generate filename
            safe_case_id = analysis_db.case_id.replace('/', '_').replace('\\', '_')
            filename = f'legal_analysis_{safe_case_id}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf'
            
            return pdf_buffer.getvalue(), filename
            
        except Exception as e:
            logger.error(f"Failed to export analysis {analysis_db.id} to PDF: {e}")
            raise
    
    def _clean_text_for_export(self, text: str) -> str:
        """Clean text for export by removing problematic characters."""
        if not text:
            return ""
        
        # Replace problematic characters
        cleaned_text = text.replace('\x00', '')  # Remove null bytes
        cleaned_text = cleaned_text.replace('\r\n', '\n')  # Normalize line endings
        cleaned_text = cleaned_text.replace('\r', '\n')
        
        return cleaned_text